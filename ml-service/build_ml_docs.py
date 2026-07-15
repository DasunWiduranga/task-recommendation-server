"""
build_ml_docs.py
================
Generates ML_PIPELINE_DOCUMENTATION.docx — a full report covering the
recommendation pipeline: dataset -> training -> output, why each model is used,
and the per-file walkthrough with full annotated source.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).parent
APP = ROOT / "app"

OUT = ROOT / "ML_PIPELINE_DOCUMENTATION.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_highlight(doc, text):
    """Highlighted main-point callout box."""
    p = doc.add_paragraph()
    run = p.add_run("KEY POINT — ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run2 = p.add_run(text)
    run2.bold = True
    run2.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    return p

def add_code_block(doc, code, lang_hint=""):
    """Insert a monospace code block (one paragraph per line)."""
    if lang_hint:
        add_para(doc, f"[{lang_hint}]", italic=True, size=9)
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()  # spacer

def read_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# Title
title = doc.add_heading("AgileAI Task Recommendation", level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sub_run = sub.add_run("Complete ML Pipeline Documentation\nDataset -> Training -> Output")
sub_run.italic = True
sub_run.font.size = Pt(13)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. Executive summary
# ---------------------------------------------------------------------------
add_heading(doc, "1. Executive Summary", level=1)
add_para(doc,
    "AgileAI is a task-recommendation engine. Given a Jira-style task "
    "(title + description) and a list of developers (with skill tags and "
    "current workload), it returns the top-5 developers most likely to "
    "complete the task successfully."
)
add_highlight(doc,
    "The system blends THREE signals into one score: how well the developer's "
    "skills match the task text (NLP), how similar developers have behaved on "
    "similar tasks in the past (Collaborative Filtering), and how much capacity "
    "the developer has left in the current sprint (workload)."
)
add_para(doc, "Final score formula:", bold=True)
add_code_block(doc,
    "final_score = W_NLP * nlp_score      (skill match)\n"
    "            + W_CF  * cf_score       (history match)\n"
    "            + W_CAP * capacity_score (workload room)\n"
    "\n"
    "Defaults: W_NLP = 0.4, W_CF = 0.4, W_CAP = 0.2\n"
    "After tuning: 0.62 / 0.18 / 0.20 (recommended)"
)

# ---------------------------------------------------------------------------
# 2. The Dataset
# ---------------------------------------------------------------------------
add_heading(doc, "2. The Dataset", level=1)
add_para(doc,
    "Training data comes from the TAWOS research corpus, a public dump of real "
    "Jira tickets (Apache Mesos, Appcelerator Titanium, etc.). The pipeline "
    "uses TWO CSV files (no SQL database is involved):"
)
add_para(doc, "Input files (folder: ./data/)", bold=True)
add_bullet(doc, "issues.csv — one row per Jira ticket: ID, Title, Description_Text, Assignee_ID, Resolution, Project_ID")
add_bullet(doc, "issue_components.csv — one row per (ticket, component) pair, used to infer each developer's skills")

add_highlight(doc,
    "We DO NOT pull data from a SQL database at training time or at request "
    "time. The training script reads CSVs; the serving API receives developers "
    "and assignments in the HTTP request body from the Node/Express backend."
)

add_para(doc, "How raw rows become training records", bold=True)
add_numbered(doc, "Keep only resolved tickets (Resolution in: Fixed, Done, Resolved, Implemented, Completed, Complete).")
add_numbered(doc, "Drop tickets without an assignee (we need (dev, task) pairs).")
add_numbered(doc, "Keep only 'active' developers — those with >= 10 resolved tickets — to avoid cold-start noise.")
add_numbered(doc, "Build a clean text column = Title + '. ' + Description_Text for NLP embedding.")
add_numbered(doc, "Infer each developer's skill tags from the components they have historically resolved (top 8 most-common components).")
add_numbered(doc, "Convert into three lists the trainer expects: developers, tasks, assignments (every resolved ticket = one positive interaction).")

# ---------------------------------------------------------------------------
# 3. The three components
# ---------------------------------------------------------------------------
add_heading(doc, "3. The Three Scoring Components", level=1)

add_heading(doc, "3.1 NLP score (skill match) — Model A", level=2)
add_para(doc,
    "Uses sentence-transformers (all-MiniLM-L6-v2, 384-dim embeddings)."
)
add_bullet(doc, "Each task description is encoded into a 384-number vector ('meaning fingerprint').")
add_bullet(doc, "Each unique developer skill word is encoded into a 384-number vector.")
add_bullet(doc, "For a (task, developer) pair, cosine-similarity is computed between the task vector and each of the developer's skill vectors.")
add_bullet(doc, "Final NLP score = average of the TOP-3 strongest skill similarities (the 'top-k trick' — rewards specialists over generalists).")
add_highlight(doc,
    "Why NLP? Because keywords lie. 'Build a containerized service' should match "
    "the skill tag 'Docker' even though the words never overlap. The embedding "
    "captures meaning, not exact tokens."
)

add_heading(doc, "3.2 CF score (history match) — Model B", level=2)
add_para(doc,
    "Collaborative Filtering with k-Nearest Neighbours over a Developer x Task "
    "interaction matrix."
)
add_bullet(doc, "Matrix cell = 1 if the developer accepted/resolved that task historically, 0 otherwise.")
add_bullet(doc, "kNN finds the 5 developers whose accept/reject pattern is most similar to the candidate developer.")
add_bullet(doc, "Predicted score for (dev, task) = average of how those 5 similar devs scored on this task.")
add_bullet(doc, "Cold start: if the dev or task wasn't seen during training, CF returns 0.5 (neutral).")
add_highlight(doc,
    "Why CF? Two developers with identical skill tags can perform very "
    "differently. CF captures the behavioural signal that NLP cannot: 'devs "
    "like this one have historically succeeded on tasks like this one'."
)

add_heading(doc, "3.3 Capacity score (workload room)", level=2)
add_para(doc, "Pure arithmetic — no ML.")
add_code_block(doc,
    "load_ratio    = min(1.0, current_story_points / sprint_capacity)\n"
    "capacity_score = 1.0 - load_ratio\n"
    "# 1.0 = totally free, 0.0 = overloaded"
)
add_highlight(doc,
    "Why capacity? The best-matched developer is useless if they're already "
    "swamped. This guard rail prevents the system from piling tasks onto one "
    "person."
)

# ---------------------------------------------------------------------------
# 4. Training pipeline
# ---------------------------------------------------------------------------
add_heading(doc, "4. Training Pipeline (offline)", level=1)
add_para(doc, "Command: python train_from_csv.py --data-dir ./data", bold=True)
add_numbered(doc, "Parse command-line args (--data-dir, --output-dir, --min-issues, --projects).")
add_numbered(doc, "Read issues.csv and issue_components.csv into pandas DataFrames.")
add_numbered(doc, "Filter rows: only resolved + assigned + active developers (>= 10 issues each).")
add_numbered(doc, "Build the three lists: developers (id + top-8 skills), tasks (id + text), assignments (every resolved ticket).")
add_numbered(doc, "Hand the lists to RecommenderModelTrainer.train_full_pipeline().")
add_numbered(doc, "PHASE 1 (NLP): encode every task description and every unique skill word; cache the 384-dim vectors.")
add_numbered(doc, "PHASE 2 (CF): build the Developer x Task matrix; normalise; fit kNN on the rows.")
add_numbered(doc, "PHASE 3 (persist): pickle both models to ./models/recommender_nlp_vlatest.pkl and recommender_cf_vlatest.pkl; write training stats to recommender_metadata.json.")

# ---------------------------------------------------------------------------
# 5. Runtime pipeline
# ---------------------------------------------------------------------------
add_heading(doc, "5. Runtime Pipeline (serving)", level=1)
add_para(doc, "Command: uvicorn app.main:app --port 8000", bold=True)
add_numbered(doc, "On startup, main.py loads the sentence-transformer base model AND the two .pkl files via initialize_models().")
add_numbered(doc, "Express backend POSTs to /recommend with: { task, developers[], assignments[], workloads{}, sprintCapacity }.")
add_numbered(doc, "main.py forwards the payload to recommender_v2.get_recommendations().")
add_numbered(doc, "For each developer in the request, compute nlp_score, cf_score, capacity_score.")
add_numbered(doc, "Blend the three with the weighted formula and clamp to [0, 1].")
add_numbered(doc, "Sort by score and return the top-5 as JSON, including a score breakdown and a human-readable reasoning string.")
add_numbered(doc, "When the user clicks accept/reject, the frontend posts to /feedback, which immediately updates the live in-memory CF matrix (no restart).")
add_numbered(doc, "Sprint-end: admin (or a scheduler) POSTs to /retrain with the full payload from MongoDB; both models are rebuilt and hot-swapped.")

# ---------------------------------------------------------------------------
# 6. Output
# ---------------------------------------------------------------------------
add_heading(doc, "6. Output Format", level=1)
add_para(doc, "Sample response from POST /recommend:", bold=True)
add_code_block(doc, """{
  "recommendations": [
    {
      "developerId": "dev_42",
      "name": "Alice Chen",
      "score": 0.812,
      "breakdown": { "nlp": 0.91, "cf": 0.78, "capacity": 0.65 },
      "skillTags": ["Python", "Docker", "PostgreSQL"],
      "reasoning": "Recommended: strong skill match, high acceptance history, good availability",
      "cold_start": false
    },
    ... up to 5 entries ...
  ],
  "cold_start": false,
  "models_loaded": true,
  "timestamp": "2026-06-27T10:24:15.123456"
}""", lang_hint="JSON")

# ---------------------------------------------------------------------------
# 7. Why use ML at all?
# ---------------------------------------------------------------------------
add_heading(doc, "7. Why Use These ML Models?", level=1)

add_heading(doc, "7.1 Why not just keyword matching?", level=2)
add_para(doc,
    "A keyword filter ('task contains \"docker\" -> developer has \"docker\" skill') "
    "is brittle. Tasks rarely state the technology by name. An issue titled "
    "'Stand up the new microservice deployment pipeline' should match Docker / "
    "Kubernetes / CI experts, but the keywords never appear."
)
add_highlight(doc, "Sentence-transformers capture MEANING, not tokens. This is what NLP buys you.")

add_heading(doc, "7.2 Why not just NLP?", level=2)
add_para(doc,
    "Two developers with the same skill tags can still be wildly different. "
    "Maybe one consistently accepts and finishes auth tickets while the other "
    "rejects them. Skill tags are static; behaviour is dynamic."
)
add_highlight(doc, "Collaborative filtering captures the behavioural pattern. It learns from outcomes, not declarations.")

add_heading(doc, "7.3 Why not just CF?", level=2)
add_para(doc,
    "CF needs history. A brand-new developer or a brand-new task has none -> "
    "cold-start problem. NLP works from day one because it only needs the "
    "developer's listed skills and the task's description."
)
add_highlight(doc, "NLP is the cold-start safety net. CF takes over once enough feedback accumulates.")

add_heading(doc, "7.4 Why include capacity?", level=2)
add_para(doc,
    "Pure ML scores ignore reality: the top-ranked developer may already be "
    "at 110% of sprint capacity. The capacity term is a hard guardrail that "
    "demotes overloaded developers so the recommendation is actionable, not "
    "theoretical."
)

add_heading(doc, "7.5 Main takeaways", level=2)
add_bullet(doc, "NLP handles cold-start, semantic matching, and brand-new tasks.")
add_bullet(doc, "CF handles personalisation — 'devs like you, on tasks like this'.")
add_bullet(doc, "Capacity ensures the recommendation is operationally feasible.")
add_bullet(doc, "All three blend into ONE score so the frontend just sorts a list.")
add_bullet(doc, "The fusion weights (W_NLP / W_CF / W_CAP) are env-vars so you can re-tune without re-deploying.")

# ---------------------------------------------------------------------------
# 8. File-by-file reference
# ---------------------------------------------------------------------------
add_heading(doc, "8. File-by-File Reference (Why + Full Annotated Code)", level=1)

# helper to render a file section
def file_section(doc, *, name, role, why_use, importance, path):
    add_heading(doc, name, level=2)
    add_para(doc, "Role:", bold=True)
    add_para(doc, role)
    add_para(doc, "Why we use this file:", bold=True)
    add_para(doc, why_use)
    add_highlight(doc, importance)
    add_para(doc, "Full annotated source:", bold=True)
    add_code_block(doc, read_file(path), lang_hint="Python")

# 8.1 train_from_csv.py
file_section(
    doc,
    name="8.1  train_from_csv.py  (ENTRYPOINT — training)",
    role=(
        "The script you run from the terminal to TRAIN the models. "
        "Reads the two CSVs, cleans the data, and hands three Python lists "
        "(developers, tasks, assignments) to RecommenderModelTrainer."
    ),
    why_use=(
        "It is the single, reproducible command that turns raw research data "
        "into the .pkl files the API loads at startup. Without it, the system "
        "has nothing to score against."
    ),
    importance=(
        "This is the only canonical training entrypoint in the project. "
        "The legacy train.py (MySQL/Kaggle path) can be deleted."
    ),
    path=ROOT / "train_from_csv.py",
)

# 8.2 app/model_trainer.py
file_section(
    doc,
    name="8.2  app/model_trainer.py  (THE BRAIN)",
    role=(
        "Contains the two ML models and the code that saves/loads them as "
        ".pkl files. Four classes: ModelPersistence, EnhancedNLPModel, "
        "EnhancedCollabFilter, RecommenderModelTrainer."
    ),
    why_use=(
        "This is where the actual machine learning happens. Encapsulates "
        "all training logic so the entrypoint script stays tiny and so the "
        "serving code can re-use the SAME model classes at request time."
    ),
    importance=(
        "Single source of truth for both models. Anything you want to change "
        "about HOW the models learn lives in this file."
    ),
    path=APP / "model_trainer.py",
)

# 8.3 app/nlp_matcher.py
file_section(
    doc,
    name="8.3  app/nlp_matcher.py  (NLP HELPER)",
    role=(
        "Loads the all-MiniLM-L6-v2 sentence-transformer once, exposes "
        "helpers to encode text into 384-dim vectors and to compute the "
        "top-k cosine-similarity score used by both the trainer and the "
        "fallback runtime path."
    ),
    why_use=(
        "Centralises the 'turn text into numbers' logic so the trainer and "
        "the live recommender share the EXACT same scoring math. Production "
        "scores therefore match the offline evaluation numbers."
    ),
    importance=(
        "The 'top-3 strongest matches' aggregation rule lives here. Change "
        "TOPN_NLP and both training and serving update together."
    ),
    path=APP / "nlp_matcher.py",
)

# 8.4 app/collab_filter.py
file_section(
    doc,
    name="8.4  app/collab_filter.py  (LIVE CF — feedback loop)",
    role=(
        "A second CF implementation that lives in memory and updates "
        "INSTANTLY when the user clicks accept/reject. Separate from the "
        "pickled CF in model_trainer.py."
    ),
    why_use=(
        "The pickled CF only updates when you re-train. This in-memory CF "
        "lets the recommender react within ONE request to fresh user "
        "feedback — important for the live demo and for retraining between "
        "sprints."
    ),
    importance=(
        "Two CF models exist on purpose: the pickled one is the slow, "
        "evaluated baseline; this one is the fast, hot-updating overlay."
    ),
    path=APP / "collab_filter.py",
)

# 8.5 app/recommender_v2.py
file_section(
    doc,
    name="8.5  app/recommender_v2.py  (REQUEST-TIME SCORER)",
    role=(
        "Called by every /recommend HTTP request. Loads the .pkl files at "
        "startup, then for each developer computes nlp_score + cf_score + "
        "capacity_score, blends them, sorts, and returns the top 5."
    ),
    why_use=(
        "It is the live decision-maker. Everything else (training, "
        "feedback, retraining) exists to make this file's output more "
        "accurate."
    ),
    importance=(
        "The fusion weights W_NLP / W_CF / W_CAP are defined here as env-"
        "vars — you tune them without re-deploying code."
    ),
    path=APP / "recommender_v2.py",
)

# 8.6 app/feedback.py
file_section(
    doc,
    name="8.6  app/feedback.py  (LIVE LEARNING LOOP)",
    role=(
        "Receives accept/reject clicks from the frontend and pushes them "
        "straight into the live in-memory CF matrix."
    ),
    why_use=(
        "Closes the feedback loop in real time. Also keeps running accept/"
        "reject counters that feed the /accuracy endpoint, so you can show "
        "model performance on a dashboard."
    ),
    importance=(
        "Every click teaches the model. Without this file, /recommend would "
        "be a static one-shot scorer."
    ),
    path=APP / "feedback.py",
)

# 8.7 app/retrain.py
file_section(
    doc,
    name="8.7  app/retrain.py  (SPRINT-END FULL REBUILD)",
    role=(
        "Rebuilds BOTH models from scratch using a full payload (real "
        "developers, sprint tasks, accept/reject history) and hot-swaps the "
        "new models into the running service without restarting."
    ),
    why_use=(
        "Lets the team replace the TAWOS-trained baseline with models "
        "trained on YOUR real history — periodically and safely (zero-"
        "downtime swap)."
    ),
    importance=(
        "Triggered by POST /retrain. Has a graceful fallback (rebuild_from_"
        "accumulated) that refreshes ONLY the CF from feedback the service "
        "has already seen."
    ),
    path=APP / "retrain.py",
)

# 8.8 app/main.py
file_section(
    doc,
    name="8.8  app/main.py  (THE WEB SERVER)",
    role=(
        "FastAPI app that wires all the modules above into HTTP endpoints: "
        "/health, /status, /recommend, /feedback, /accuracy, /cf/stats, "
        "/retrain, /train/status."
    ),
    why_use=(
        "The single external contract. The React + Express front-of-house "
        "talks ONLY to these endpoints — no shared database, no shared "
        "imports. That's why this service has zero SQL dependencies."
    ),
    importance=(
        "On startup it (1) loads the sentence-transformer, (2) loads the "
        ".pkl files, (3) restores feedback from disk. Then it's ready."
    ),
    path=APP / "main.py",
)

# 8.9 app/schemas.py
file_section(
    doc,
    name="8.9  app/schemas.py  (REQUEST/RESPONSE SHAPES)",
    role=(
        "Pydantic models that FastAPI uses to validate every request body "
        "and serialise every response."
    ),
    why_use=(
        "Forces a strict contract with the Express backend. Bad payloads "
        "are rejected at the door (HTTP 422) instead of crashing the "
        "scorer. Doubles as documentation for the API."
    ),
    importance=(
        "If you add a new field to /recommend or /feedback, you change it "
        "here — and FastAPI updates the OpenAPI docs automatically."
    ),
    path=APP / "schemas.py",
)

# ---------------------------------------------------------------------------
# Closing
# ---------------------------------------------------------------------------
add_heading(doc, "9. Summary Table", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "File"
hdr[1].text = "Phase"
hdr[2].text = "Why it exists"

rows = [
    ("train_from_csv.py",     "Training entry",  "Turn CSVs into .pkl files"),
    ("app/model_trainer.py",  "Training core",   "Holds the NLP + CF model classes"),
    ("app/nlp_matcher.py",    "Shared NLP",      "Text -> 384-dim vectors; top-k score"),
    ("app/collab_filter.py",  "Live CF",         "In-memory CF that learns on every click"),
    ("app/recommender_v2.py", "Serving core",    "Computes the final blended score"),
    ("app/feedback.py",       "Feedback loop",   "Receives accept/reject, updates live CF"),
    ("app/retrain.py",        "Sprint refresh",  "Full rebuild + zero-downtime hot swap"),
    ("app/main.py",           "HTTP server",     "FastAPI app exposing the endpoints"),
    ("app/schemas.py",        "API contract",    "Pydantic request/response shapes"),
]
for f, phase, why in rows:
    row = table.add_row().cells
    row[0].text = f
    row[1].text = phase
    row[2].text = why

doc.add_paragraph()
add_para(doc,
    "End of documentation. Generated automatically from the source — keep this "
    "file in sync by re-running build_ml_docs.py whenever the code changes.",
    italic=True, size=9
)

doc.save(OUT)
print(f"Wrote: {OUT}")
